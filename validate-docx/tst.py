from docx import Document

def count_tables(docx_path):
    # відкриваємо документ
    doc = Document(docx_path)
    # отримуємо кількість таблиць
    return len(doc.tables)

def iterate(docx_path):
    doc = Document(docx_path)
    for i, table in enumerate(doc.tables, start=1):
        print(f"Таблиця {i}:")
        print(f"  Стиль: {table.style}")
        print(f"  Рядків: {len(table.rows)}")
        # print(f"  Стовпців: {len(table.columns)}")
        # print(f"  Автопідгін: {table.allow_autofit}")
        # for r, row in enumerate(table.rows):
        #     for c, cell in enumerate(row.cells):
        #         print(f"    [{r},{c}] -> {cell.text}")

def parse(path):
    doc = Document(path)
    table = doc.tables[2]

    print(f"  Рядків: {len(table.rows)}")

if __name__ == "__main__":
    path = "car2.docx"
    parse(path)

# print(f"  Стиль: {table.style}")