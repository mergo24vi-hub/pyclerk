import os
from docx import Document

def get_effective_font_name(run, paragraph):
    """Повертає фактичне ім'я шрифта для run, враховуючи спадкування стилів."""
    if run.font.name:
        return run.font.name
    elif run.style and run.style.font and run.style.font.name:
        return run.style.font.name
    elif paragraph.style and paragraph.style.font and paragraph.style.font.name:
        return paragraph.style.font.name
    else:
        return None

def find_non_tnr_words_with_context(doc_path):
    results = []

    def extract_words_with_fonts(paragraphs, context, file_name):
        for para in paragraphs:
            words_fonts = []
            for run in para.runs:
                font = get_effective_font_name(run, para)
                words = run.text.split()
                for word in words:
                    words_fonts.append((word, font))
            for i, (word, font) in enumerate(words_fonts):
                if font and font != "Times New Roman":
                    before = words_fonts[i - 1][0] if i > 0 else "<початок>"
                    after = words_fonts[i + 1][0] if i + 1 < len(words_fonts) else "<кінець>"
                    results.append((file_name, context, before, word, after, font))

    try:
        doc = Document(doc_path)
        extract_words_with_fonts(doc.paragraphs, "Абзац", os.path.basename(doc_path))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract_words_with_fonts(cell.paragraphs, "Таблиця", os.path.basename(doc_path))
    except Exception as e:
        print(f"⚠️ Помилка при обробці '{doc_path}': {e}")

    return results

def process_all_docs_in_subfolders(base_folder):
    all_results = []
    for entry in os.listdir(base_folder):
        subfolder_path = os.path.join(base_folder, entry)
        if os.path.isdir(subfolder_path):
            print(subfolder_path)
            for file in os.listdir(subfolder_path):
                if file.endswith(".docx") and not file.startswith("~$"):  # ігнорування тимчасових файлів Word
                    full_path = os.path.join(subfolder_path, file)
                    res = find_non_tnr_words_with_context(full_path)
                    all_results.extend(res)
    return all_results

# 🔽 Задати шлях до головної папки
root_dir = r'G:\.shortcut-targets-by-id\1Pcnp8gnqT8NS3Zl5AOanpcBmZLHuuv5I\РО робоча\\РОБОЧА ГОДЗЕВИЧ\Списання ЗЧ\\'

base_folder = r'G:\.shortcut-targets-by-id\1Pcnp8gnqT8NS3Zl5AOanpcBmZLHuuv5I\РО робоча\\РОБОЧА ГОДЗЕВИЧ\Списання ЗЧ\\'

# 🔍 Обробка
results = process_all_docs_in_subfolders(base_folder)

# 📄 Вивід
if results:
    for file, context, before, word, after, font in results:
        print(f"[{file} | {context}] ... {before} >>{word}<< {after} (шрифт: {font})")

