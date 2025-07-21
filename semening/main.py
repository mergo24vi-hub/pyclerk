#G:\.shortcut-targets-by-id\1Pcnp8gnqT8NS3Zl5AOanpcBmZLHuuv5I\РО робоча\ОСОБИСТІ ПАПКИ
import os
from docx import Document
import re
from doc2docx import convert_doc_to_docx
root_dir = r'G:\.shortcut-targets-by-id\1Pcnp8gnqT8NS3Zl5AOanpcBmZLHuuv5I\РО робоча\\РОБОЧА ГОДЗЕВИЧ\семенезація\\'


def process_document(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        for i in range(len(words) - 1):
            # Шукаємо шаблон: слово з великої літери + слово у верхньому регістрі
            if (words[i][0].isupper() and words[i][1:].islower() and
                words[i + 1].isupper()):
                # Формуємо ПІБ
                name = f"{words[i]} {words[i + 1]}"

                # Отримуємо до 7 слів перед ПІБ
                start_idx = max(0, i - 7)
                context_before = words[start_idx:i]
                context_text = " ".join(context_before)

                # Виводимо результат
                print("Контекст:", context_text)
                print("ПІБ:", name)
                # new_name = "Новий ПрізвищеІм’я"
                # paragraph.text = paragraph.text.replace(old_name, new_name)
                break

    # doc.save(file_path)  # Зберігання зміненого документа (за потреби)

def extract_all_text(doc_path):
    doc = Document(doc_path)
    full_text = []

    # абзаци документа
    for para in doc.paragraphs:
        full_text.append(para.text)

    # текст із таблиць
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    # об’єднуємо в один рядок для пошуку
    return "\n".join(full_text)

def check_signature_pattern(text):
    # Регулярний вираз:
    # .*? — лінйний пошук між частинами
    pattern = re.compile(
        r"Командир\s+військової\s+частини\s+А4007.*?"
        r"полковник.*?"
        r"Валерій\s+СЕМЕНЕЦЬ",
        re.DOTALL | re.IGNORECASE
    )
    return bool(pattern.search(text))

i=0
for folder in os.listdir(root_dir):
    folder_path = os.path.join(root_dir, folder)
    if not os.path.isdir(folder_path):
        continue

    print(folder_path)
    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)
        if not os.path.isfile(file_path):
            continue
        #print(file_path)
        # Перевіряємо, що це файл .doc з 'сдд' у назві (незалежно від регістру)
        if not file.startswith('~') and (file.lower().endswith('.docx')):

            text = extract_all_text(file_path)
            if not check_signature_pattern(text):
                print(file_path)
                print("Збіг НЕ знайдено ❌")







    i=i+1
    if i>500:
       break
    print('--------')