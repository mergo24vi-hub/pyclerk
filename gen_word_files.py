import pandas as pd
from docx import Document
from docx.shared import Cm
import os
from pathlib import Path
from datetime import datetime

# === Налаштування ===
excel_file = 'personal/шаблон.сзч.xlsx'  # шлях до твого Excel-файлу
template_file = 'personal/imaged_template.docx'
output_dir = 'generated'
#output_dir = r'G:\.shortcut-targets-by-id\1Pcnp8gnqT8NS3Zl5AOanpcBmZLHuuv5I\РО робоча\ОСОБИСТІ ПАПКИ\.Не штатні\.СЗЧ\\'
#os.makedirs(output_dir, exist_ok=True)

# === Зчитування Excel ===
df = pd.read_excel(excel_file, nrows=1)
df = df.fillna('')

def format_cell(value):
    if pd.isna(value):
        return ''
    elif isinstance(value, (datetime, pd.Timestamp)):
        return value.strftime('%d.%m.%Y')  # або '%Y-%m-%d' для ISO-формату
    return str(value)

def replace_text(paragraphs, replacements):
    for paragraph in paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, str(val))

def insert_image_in_place_of_marker(doc, marker, image_path, width_cm=5, height_cm=5):
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
            return True
    return False

# === Обробка кожного рядка ===
for index, row in df.iterrows():
    # Відкриваємо шаблон
    doc = Document(template_file)

    # Словник тегів і значень
    replacements = {
        '«Звання»': row.get('Звання', ''),
        '«ПІБ»': row.get('ПІБ', ''),
        '«Посада»': row.get('Посада', ''),
        '«ДатаНар»': format_cell(row.get('ДатаНар', '')),
        '«МісцеНар»': row.get('МісцеНар', ''),
        '«Призов»': row.get('Призов', ''),
        '«СімейнийСтан»': row.get('СімейнийСтан', ''),
        '«Дружина»': row.get('Дружина', ''),
        '«Діти»': row.get('Діти', ''),
        '«Батько»': row.get('Батько', ''),
        '«Мати»': row.get('Мати', ''),
        '«Освіта»': row.get('Освіта', ''),
        '«Адреса»': row.get('Адреса', ''),
        '«Телефон»': str(row.get('Телефон', '')),
        '«ІПН»': str(row.get('ІПН', '')),
        '«ГрупаКрові»': row.get('ГрупаКрові', '')
    }

    replace_text(doc.paragraphs, replacements)

    # Заміна в таблицях (якщо є)

    for table in doc.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for paragraph in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, str(val))

        # Вставка зображення

    image_path = 'personal//'+ row.get('Фото', '').strip()
    print(image_path)
    if image_path and os.path.isfile(image_path):
        print('photo exists')
        inserted = insert_image_in_place_of_marker(doc, 'ФОТО', image_path)
        if not inserted:
            print(f"⚠️ Не знайдено маркер 'ФОТО' у шаблоні для: {row.get('ПІБ')}")

    # Формуємо ім’я файлу
    pib_safe = str(row.get('ПІБ', 'невідомий')).replace(' ', '_')
    filename = f"СДД_{pib_safe}_сзч.docx"
    path = (os.path.join(output_dir, row.get('Папка', 'ххх'))); print(path)
    folder = Path(path)
    for file in folder.iterdir():
        if (file.is_file() and file.name.startswith('сдд') and ('стара' not in file.name)
                and os.path.splitext(file.name)[1].lower() == ".docx"):
            print(f"Знайдено файл: {file.name}")
            file.unlink()  # видалення файла
            break  # зупиняємось після першого знайденого
    doc.save(os.path.join(path, filename))

print("✅ Генерація завершена!")
